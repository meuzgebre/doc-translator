from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from doctranslator.utils.file_parser import get_file_name_with_language_extension
from doctranslator.utils.translation_api import TranslationAPI

class TranslationService:
    def __init__(self):
        self.translation_api = TranslationAPI()

    def set_api_key(self, api_key):
        self.translation_api.set_api_key(api_key)

    def translate_document(self, file_path, api_key, target_language, source_language, engine, api_template):
        doc = Document(file_path)

        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Header") or paragraph.style.name.startswith("Footer"):
                continue

            translated_text = self.translation_api.translate_text(paragraph.text, target_language, source_language)
            paragraph.text = translated_text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    translated_text = self.translation_api.translate_text(cell.text, target_language)
                    cell.text = translated_text

        translated_file_path = get_file_name_with_language_extension(file_path, target_language)
        doc.save(translated_file_path)

        return translated_file_path
    
# Todo: Select API engine
